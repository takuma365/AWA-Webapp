import os
import re
import uuid
import tempfile
import zipfile
import shutil
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from lxml import etree
from docx import Document
from django.conf import settings


class DocumentParser:
    """Wordドキュメントをパースするクラス"""
    
    def __init__(self, conversion_setting):
        """
        初期化
        
        Args:
            conversion_setting: 変換設定
        """
        self.setting = conversion_setting
        self.css_class_prefix = self.setting.css_class_prefix or ""
    
    def parse_document(self, doc_file) -> Dict[str, Any]:
        """
        Wordファイルをパースする
        
        Args:
            doc_file: Wordファイルオブジェクト
            
        Returns:
            Dict: パースされたドキュメント構造
        """
        # 一時ファイルに保存
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            for chunk in doc_file.chunks():
                tmp_file.write(chunk)
        
        xml_files = self.extract_and_save_xml(tmp_file.name, doc_file.name)
                
        try:
            doc = Document(tmp_file.name)
            
            result = {
                "paragraphs": [],
                "tables": [],
                "images": [],
                "file_path": tmp_file.name,  # 画像抽出用
                "xml_files": xml_files,  # 抽出されたXMLファイルのパス情報
            }
            
            # 段落の処理
            for paragraph in doc.paragraphs:
                parsed_paragraph = self._parse_paragraph(paragraph)
                if parsed_paragraph:
                    result["paragraphs"].append(parsed_paragraph)
            
            # テーブルの処理
            for table in doc.tables:
                parsed_table = self._parse_table(table)
                if parsed_table:
                    result["tables"].append(parsed_table)
                    
            # 画像の処理（リファレンスのみ保存）
            if self.setting.preserve_images:
                result["images"] = self._extract_image_references(doc)
            
            return result
            
        except Exception as e:
            raise Exception(f"ドキュメントの解析中にエラーが発生しました: {str(e)}")
        finally:
            # 一時ファイルを削除
            if os.path.exists(tmp_file.name):
                os.unlink(tmp_file.name)
    
    def extract_and_save_xml(self, docx_path, original_filename) -> Dict[str, str]:
        """
        Wordファイル(docx)からXMLファイルを抽出して保存する
        
        Args:
            docx_path: docxファイルのパス
            original_filename: 元のファイル名
            
        Returns:
            Dict: 保存されたXMLファイルの情報
        """
        # 保存先ディレクトリの設定
        xml_dir_base = os.path.join(settings.BASE_DIR, 'xml_data')
        if not os.path.exists(xml_dir_base):
            os.makedirs(xml_dir_base)
        
        # 一意のディレクトリ名を生成（オリジナルファイル名を含む）
        safe_filename = re.sub(r'[^\w\-\.]', '_', os.path.splitext(original_filename)[0])
        xml_dir = os.path.join(xml_dir_base, f"{safe_filename}_{uuid.uuid4().hex[:8]}")
        os.makedirs(xml_dir)
        
        saved_files = {}
        
        # docxファイルをZIPとして解凍
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            # 必要なXMLファイルを抽出
            for file_info in zip_ref.infolist():
                if file_info.filename.endswith('.xml'):
                    # XMLファイルの取得
                    xml_content = zip_ref.read(file_info.filename)
                    
                    # ファイル名の生成
                    filename = os.path.basename(file_info.filename)
                    save_path = os.path.join(xml_dir, filename)
                    
                    # XMLファイルを保存
                    with open(save_path, 'wb') as f:
                        f.write(xml_content)
                    
                    # パスをマップに追加
                    saved_files[filename] = save_path
                    
                    # document.xmlとstyles.xmlの場合は整形して保存
                    if filename in ['document.xml', 'styles.xml']:
                        try:
                            # XMLを整形
                            parser = etree.XMLParser(remove_blank_text=True)
                            tree = etree.fromstring(xml_content, parser)
                            pretty_xml = etree.tostring(tree, pretty_print=True, encoding='UTF-8').decode('utf-8')
                            
                            # 整形されたXMLを保存
                            pretty_path = os.path.join(xml_dir, f"pretty_{filename}")
                            with open(pretty_path, 'w', encoding='utf-8') as f:
                                f.write(pretty_xml)
                            
                            saved_files[f"pretty_{filename}"] = pretty_path
                        except Exception as e:
                            print(f"XMLの整形中にエラーが発生しました: {str(e)}")
        
        # docxの構造をわかりやすくするためのディレクトリ構造も保存
        structure_path = os.path.join(xml_dir, 'docx_structure.txt')
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            with open(structure_path, 'w', encoding='utf-8') as f:
                f.write("# Docxファイル内の構造\n\n")
                for file_info in sorted(zip_ref.infolist(), key=lambda x: x.filename):
                    f.write(f"{file_info.filename} - {file_info.file_size} bytes\n")
        
        saved_files['structure'] = structure_path
        
        return {
            'directory': xml_dir,
            'files': saved_files
        }
    
    def _parse_paragraph(self, paragraph) -> Optional[Dict[str, Any]]:
        """段落をパースする"""
        # テキストがNoneの場合を処理
        text = paragraph.text or ""
        
        # 空の段落を削除する設定がオンで、テキストが空なら無視
        if not text.strip() and self.setting.remove_empty_paragraphs:
            return None
        
        result = {
            "text": text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "runs": []
        }
        
        # 段落内のテキスト処理
        for run in paragraph.runs:
            parsed_run = self._parse_run(run)
            if parsed_run:
                result["runs"].append(parsed_run)
        
        return result
    
    def _parse_run(self, run) -> Optional[Dict[str, Any]]:
        """テキストスタイルをパースする"""
        # テキストがNoneの場合を処理
        text = run.text or ""
        
        if not text.strip():
            return None
        
        # 色の取得（エラーハンドリング付き）
        color = None
        try:
            if run.font.color and run.font.color.rgb:
                color = '#{:06x}'.format(int(str(run.font.color.rgb), 16))
        except Exception:
            pass
        
        return {
            "text": text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "color": color
        }
    
    def _parse_table(self, table) -> Dict[str, Any]:
        """テーブルをパースする"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    parsed_paragraph = self._parse_paragraph(paragraph)
                    if parsed_paragraph:
                        cell_content.append(parsed_paragraph)
                cells.append({"content": cell_content})
            rows.append({"cells": cells})
        
        return {
            "rows": rows
        }
    
    def _extract_image_references(self, doc) -> List[Dict[str, Any]]:
        """ドキュメントから画像参照を抽出する"""
        image_refs = []
        
        try:
            # 画像を含むリレーションシップを探す
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    # 拡張子の取得
                    ext = os.path.splitext(rel.target_ref)[1].lower()
                    
                    # 一意のファイル名を生成
                    filename = f"{uuid.uuid4()}{ext}"
                    
                    image_refs.append({
                        "path": rel.target_ref,
                        "target_part": rel.target_part,
                        "filename": filename,
                        "ext": ext
                    })
        except Exception as e:
            print(f"画像参照の抽出中にエラーが発生しました: {str(e)}")
        
        return image_refs


class ImageHandler:
    """画像を処理するクラス"""
    
    def __init__(self, conversion_setting):
        """
        初期化
        
        Args:
            conversion_setting: 変換設定
        """
        self.setting = conversion_setting
        self.images = []
        
        # 画像保存先を設定
        image_dir = self.setting.image_dir or 'images'
        self.media_root = Path(settings.UPLOAD_IMAGE_DIR)
        
        # 画像保存ディレクトリが存在しなければ作成
        if not self.media_root.exists():
            self.media_root.mkdir(parents=True, exist_ok=True)
    
    def extract_and_save_images(self, image_refs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        画像を抽出して保存する
        
        Args:
            image_refs: 画像参照リスト
            
        Returns:
            保存された画像情報のリスト
        """
        saved_images = []
        
        for image_ref in image_refs:
            try:
                # 画像データを取得
                image_data = image_ref["target_part"].blob
                
                # 保存先のパス設定
                filename = image_ref["filename"]
                output_path = self.media_root / filename
                image_dir = self.setting.image_dir or 'images'
                relative_path = f"{image_dir}/{filename}"
                
                # 画像の保存
                with open(output_path, "wb") as img_file:
                    img_file.write(image_data)
                
                # 画像情報を追加
                saved_image = {
                    "path": str(output_path),
                    "url": f"{settings.MEDIA_URL}{relative_path}",
                    "filename": filename
                }
                saved_images.append(saved_image)
                self.images.append(saved_image)
            
            except Exception as e:
                print(f"画像 {image_ref.get('filename', '不明')} の保存中にエラーが発生しました: {str(e)}")
        
        return saved_images


class HTMLGenerator:
    """HTMLを生成するクラス"""
    
    def __init__(self, conversion_setting):
        """
        初期化
        
        Args:
            conversion_setting: 変換設定
        """
        self.setting = conversion_setting
        # css_class_prefixは残しておくが使用しない
        self.css_class_prefix = self.setting.css_class_prefix or ""
    
    def generate_html(self, parsed_data: Dict[str, Any], saved_images: List[Dict[str, Any]] = None) -> str:
        """
        パースされたデータからHTMLを生成する
        
        Args:
            parsed_data: パースされたデータ
            saved_images: 保存された画像情報
            
        Returns:
            HTML文字列
        """
        html_parts = []
        
        # HTMLのルート要素 - 接頭辞を使用しない
        html_root = '<div class="document">'
        html_parts.append(html_root)
        
        # 段落の処理
        for paragraph in parsed_data.get("paragraphs", []):
            html_parts.append(self._convert_paragraph(paragraph))
        
        # テーブルの処理
        for table in parsed_data.get("tables", []):
            html_parts.append(self._convert_table(table))
        
        # 画像の処理
        if saved_images:
            for image in saved_images:
                html_parts.append(self._insert_image(image))
        
        html_parts.append('</div>')
        
        return "\n".join(html_parts)
    
    def _convert_paragraph(self, paragraph: Dict[str, Any]) -> str:
        """段落をHTMLに変換"""
        style_name = paragraph.get("style", "Normal").lower().replace(" ", "-")
        
        # 見出しの場合
        if paragraph.get("style", "").startswith("Heading"):
            level = paragraph.get("style", "")[-1] if paragraph.get("style", "")[-1:].isdigit() else "3"
            tag = f"h{level}"
        else:
            tag = "p"
        
        # CSSクラスを設定 - 接頭辞を使用しない
        css_class = f"paragraph {style_name}"
        
        # テキストのスタイル処理
        runs = paragraph.get("runs", [])
        if runs:
            content = []
            for run in runs:
                content.append(self._convert_run(run))
            inner_html = "".join(content)
        else:
            inner_html = paragraph.get("text", "")
        
        return f'<{tag} class="{css_class}">{inner_html}</{tag}>'
    
    def _convert_run(self, run: Dict[str, Any]) -> str:
        """テキストスタイルをHTMLに変換"""
        text = run.get("text", "")
        
        # スタイルの適用
        if run.get("bold"):
            text = f'<strong>{text}</strong>'
        if run.get("italic"):
            text = f'<em>{text}</em>'
        if run.get("underline"):
            text = f'<u>{text}</u>'
        
        # 色の処理
        color = run.get("color")
        if color:
            text = f'<span style="color: {color}">{text}</span>'
        
        return text
    
    def _convert_table(self, table: Dict[str, Any]) -> str:
        """テーブルをHTMLに変換"""
        html_parts = ['<table class="table">']
        html_parts.append('<tbody>')
        
        for row in table.get("rows", []):
            html_parts.append('<tr class="row">')
            for cell in row.get("cells", []):
                cell_content = []
                for paragraph in cell.get("content", []):
                    cell_content.append(self._convert_paragraph(paragraph))
                html_parts.append(f'<td class="cell">{"".join(cell_content)}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</tbody>')
        html_parts.append('</table>')
        
        return "\n".join(html_parts)
    
    def _insert_image(self, image: Dict[str, Any]) -> str:
        """画像をHTMLに挿入"""
        return f'<p><img class="image" src="{image["url"]}" alt="" /></p>'


class RuleApplier:
    """変換ルールを適用するクラス"""
    
    def __init__(self, conversion_setting):
        """
        初期化
        
        Args:
            conversion_setting: 変換設定
        """
        self.setting = conversion_setting
    
    def apply_rules(self, html_content: str) -> str:
        """
        変換ルールを適用する
        
        Args:
            html_content: HTML文字列
            
        Returns:
            ルール適用後のHTML文字列
        """
        from lxml import etree
        
        # HTML解析
        parser = etree.HTMLParser()
        try:
            html_tree = etree.fromstring(f"<div>{html_content}</div>", parser)
        except Exception as e:
            print(f"HTMLの解析中にエラーが発生しました: {str(e)}")
            return html_content
        
        # ルールの適用（優先度順）
        for rule in self.setting.rules.filter(active=True).order_by('priority'):
            try:
                # ルールタイプごとの処理
                if rule.rule_type == 'tag_replace':
                    self._apply_tag_replace(html_tree, rule)
                elif rule.rule_type == 'class_add':
                    self._apply_class_add(html_tree, rule)
                elif rule.rule_type == 'attribute_add':
                    self._apply_attribute_add(html_tree, rule)
                elif rule.rule_type == 'custom':
                    self._apply_custom_rule(html_tree, rule)
            except Exception as e:
                print(f"ルール '{rule.name}' の適用中にエラーが発生しました: {str(e)}")
        
        # HTMLに変換
        result = etree.tostring(html_tree, encoding='unicode', pretty_print=True, method="html")
        
        # 最初のdivを除去して返す
        match = re.search(r'<div>(.*)</div>', result, re.DOTALL)
        if match:
            return match.group(1)
        return result
    
    def _apply_tag_replace(self, html_tree, rule):
        """タグ置換ルールの適用"""
        for elem in html_tree.xpath(rule.source_selector):
            new_elem = etree.Element(rule.target_value)
            new_elem.text = elem.text
            new_elem.tail = elem.tail
            
            # 属性のコピー
            for key, value in elem.attrib.items():
                new_elem.set(key, value)
            
            # 子要素のコピー
            for child in elem:
                new_elem.append(child)
            
            # 要素の置換
            parent = elem.getparent()
            if parent is not None:
                parent.replace(elem, new_elem)
    
    def _apply_class_add(self, html_tree, rule):
        """クラス追加ルールの適用"""
        for elem in html_tree.xpath(rule.source_selector):
            classes = elem.get('class', '').split()
            new_classes = rule.target_value.split()
            
            # 重複を避けて新しいクラスを追加
            for cls in new_classes:
                if cls not in classes:
                    classes.append(cls)
            
            elem.set('class', ' '.join(classes))
    
    def _apply_attribute_add(self, html_tree, rule):
        """属性追加ルールの適用"""
        attr_name, attr_value = rule.target_value.split('=', 1) if '=' in rule.target_value else (rule.target_value, '')
        
        # 値が引用符で囲まれている場合は取り除く
        attr_value = attr_value.strip('"\'')
        
        for elem in html_tree.xpath(rule.source_selector):
            elem.set(attr_name, attr_value)
    
    def _apply_custom_rule(self, html_tree, rule):
        """カスタム処理ルールの適用"""
        # カスタムルールの実装（必要に応じて拡張）
        pass


class WordToHtmlConverter:
    """WordファイルをHTML形式に変換するクラス"""
    
    def __init__(self, conversion_setting):
        """
        コンストラクタ
        
        Args:
            conversion_setting (ConversionSetting): 変換設定
        """
        self.setting = conversion_setting
        self.document_parser = DocumentParser(conversion_setting)
        self.image_handler = ImageHandler(conversion_setting)
        self.html_generator = HTMLGenerator(conversion_setting)
        self.rule_applier = RuleApplier(conversion_setting)
        self.parsed_data = None  # パース結果を保持するフィールドを追加
    
    def convert(self, word_file) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Wordファイルを変換してHTMLを返す
        
        Args:
            word_file: WordファイルのFileオブジェクト
            
        Returns:
            tuple: (HTML文字列, 画像ファイルリスト)
        """
        try:
            # ドキュメントのパース
            self.parsed_data = self.document_parser.parse_document(word_file)
            
            # 画像の抽出と保存
            saved_images = []
            if self.setting.preserve_images and self.parsed_data.get("images"):
                saved_images = self.image_handler.extract_and_save_images(self.parsed_data["images"])
            
            # HTMLの生成
            html_content = self.html_generator.generate_html(self.parsed_data, saved_images)
            
            # 変換ルールの適用
            if self.setting.rules.filter(active=True).exists():
                html_content = self.rule_applier.apply_rules(html_content)
            
            return html_content, self.image_handler.images
            
        except Exception as e:
            import traceback
            print(f"変換中にエラーが発生しました: {traceback.format_exc()}")
            raise Exception(f"変換中にエラーが発生しました: {str(e)}") 